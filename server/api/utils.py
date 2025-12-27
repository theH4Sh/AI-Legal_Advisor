import os
from docx import Document
from django.conf import settings
from copy import deepcopy

def replace_in_paragraph(paragraph, replacements):
    """Replace placeholders in a paragraph's runs preserving formatting."""
    # Join all runs to a single text, do replacements, then rewrite runs minimally
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    for key, val in replacements.items():
        new_text = new_text.replace("{" + key + "}", str(val))
    # If nothing changed, skip
    if new_text == full_text:
        return
    # Clear runs and add new run with new_text (this sacrifices some run-level formatting)
    # For small templates this is usually fine; for preserving styling per run you'd need more advanced logic.
    for i in range(len(paragraph.runs) - 1, -1, -1):
        paragraph.runs[i].clear()
    paragraph.add_run(new_text)

def replace_docx_placeholders(doc: Document, replacements: dict):
    # paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

    # headers / footers
    for section in doc.sections:
        header = section.header
        footer = section.footer
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph, replacements)

def generate_docx_from_template(template_path: str, replacements: dict, output_path: str):
    """
    template_path: filesystem path to a .docx template
    replacements: dict of placeholder -> value (placeholders without braces)
    output_path: path to save resulting .docx
    """
    doc = Document(template_path)
    replace_docx_placeholders(doc, replacements)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
